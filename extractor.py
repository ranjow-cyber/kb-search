# ============================================================
#  extractor.py — Ekstrakcja tekstu z plików PDF, DOCX, MD/TXT
# ============================================================

from __future__ import annotations
import re
from pathlib import Path


def extract_text(file_path: str | Path) -> str:
    """
    Wyekstrahuj tekst z pliku PDF, DOCX lub Markdown/TXT.
    Zwraca czysty tekst gotowy do indeksowania.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(path)
    elif suffix == ".docx":
        return _extract_docx(path)
    elif suffix in (".md", ".txt"):
        return _extract_text_file(path)
    else:
        raise ValueError(f"Nieobsługiwany typ pliku: {suffix}")


def _extract_pdf(path: Path) -> str:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("Zainstaluj: pip install pdfplumber")

    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())

    return _clean("\n\n".join(pages))


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Zainstaluj: pip install python-docx")

    doc = Document(path)
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return _clean("\n\n".join(paragraphs))


def _extract_text_file(path: Path) -> str:
    text = path.read_text(encoding="utf-8", errors="replace")
    text = re.sub(r"!\[.*?\]\(.*?\)", "", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    text = re.sub(r"[#*`_~>]+", "", text)
    return _clean(text)


def _clean(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def chunk_text(
    text: str,
    chunk_size: int = 500,
    overlap: int = 50,
) -> list[str]:
    """
    Podziel tekst na nakładające się fragmenty.
    """
    words = text.split()
    if not words:
        return []

    chunks: list[str] = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunks.append(" ".join(words[start:end]))
        if end == len(words):
            break
        start += chunk_size - overlap

    return chunks
